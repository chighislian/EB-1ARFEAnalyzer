import os
import json
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_COLOR_INDEX
from docx2pdf import convert

# === CONFIG ===
INPUT_FOLDER = "petition_analysis"
OUTPUT_FOLDER = "reports"
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# === Heuristic: Determine severity based on keywords ===
def get_severity(pattern):
    pattern = pattern.lower()
    if any(word in pattern for word in ["no evidence", "not provided", "failed to", "insufficient", "does not", "lack", "unable"]):
        return "high"
    elif any(word in pattern for word in ["limited", "unclear", "vague", "generic"]):
        return "medium"
    else:
        return "low"

# === Suggestions based on keywords ===
SUGGESTIONS = {
    "no evidence": "Add third-party documentation or supporting materials.",
    "not provided": "Include the missing data or supporting letter.",
    "insufficient": "Provide more detailed and verifiable documentation.",
    "vague": "Clarify claims with specific examples and data.",
    "lack": "Submit corroborating evidence from independent sources.",
    "generic": "Use field-specific metrics or endorsements to boost credibility.",
    "failed to": "Re-evaluate this section with the USCIS criteria checklist."
}

def get_suggestion(pattern):
    for keyword, suggestion in SUGGESTIONS.items():
        if keyword in pattern.lower():
            return suggestion
    return "Consider strengthening this section with clearer and more credible evidence."

def get_reviewer_note():
    return "This issue reflects a common reason USCIS may deny a case. Additional evidence or clarification is strongly advised."

def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)

def add_colored_paragraph(doc, text, severity="medium"):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(11)
    if severity == "high":
        run.font.highlight_color = WD_COLOR_INDEX.RED
    elif severity == "medium":
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    elif severity == "low":
        run.font.highlight_color = WD_COLOR_INDEX.GRAY_25

def generate_report(json_path):
    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    filename = os.path.basename(json_path).replace(".json", "")
    doc = Document()
    doc.add_heading("EB-1A Petition RFE Risk Assessment", 0)
    doc.add_paragraph(f"Filename: {data['filename']}")
    doc.add_paragraph("Generated QA Memo for Pre-Submission Review")

    add_heading(doc, "Executive Summary", level=1)
    doc.add_paragraph("This report summarizes red flags detected in the petition draft, classified under EB-1A criteria. It includes excerpt patterns, severity ratings, improvement suggestions, and simulated reviewer notes.")

    add_heading(doc, "Table of Contents", level=1)
    for section in data.get("sections", {}).keys():
        #doc.add_paragraph(f"• {section.title()}", style='List Bullet')
        doc.add_paragraph(section.title(), style='List Bullet')

    add_heading(doc, "Risk Matrix", level=1)
    for section, items in data.get("sections", {}).items():
        if not items:
            continue
        add_heading(doc, f"Section: {section.title()}", level=2)
        for item in items:
            criterion = item["criterion"]
            pattern = item["pattern"]
            sentence = item["sentence"]
            severity = get_severity(pattern)
            suggestion = get_suggestion(pattern)
            reviewer_note = get_reviewer_note()

            add_colored_paragraph(doc, f"- ❗ {criterion} → '{pattern}'", severity)
            doc.add_paragraph(f"  ↳ Found in: {sentence}")
            doc.add_paragraph(f"  ↳ Severity: {severity.capitalize()}")
            doc.add_paragraph(f"  ↳ Suggestion: {suggestion}")
            doc.add_paragraph(f"  🧠 Reviewer Note: {reviewer_note}")

    # Save DOCX
    report_path = os.path.join(OUTPUT_FOLDER, f"{filename}_qa_report.docx")
    doc.save(report_path)
    print(f"✅ DOCX report generated: {report_path}")

    # Convert to PDF
    pdf_path = report_path.replace(".docx", ".pdf")
    convert(report_path, pdf_path)
    print(f"✅ PDF version saved: {pdf_path}")

# Run on all petition_analysis JSONs
if __name__ == "__main__":
    for file in os.listdir(INPUT_FOLDER):
        if file.endswith(".json"):
            generate_report(os.path.join(INPUT_FOLDER, file))